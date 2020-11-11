from docx import Document

doc = Document('test.docx')


def how_many_tags(paragraph_text):
    n = 0
    char = 0
    while char != len(paragraph_text):
        if paragraph_text[char] == '{' and paragraph_text[char + 1] == '{':
            n += 1
        char += 1
    return n


def find_tag_in_text(number_of_tags, text, list_of_tags_from_strs):
    while number_of_tags != 0:
        start = text.find('{{')  # символы начала тега
        end = text.find('}}', start)  # символы конца тега
        if start != -1 and end != -1 and text != '':  # условия, чтобы пустые строки не попали в список
            list_of_tags_from_strs.append(text[start + 2:end])
        text = text[end + 2:]
        # переприсваивание строки, обрезая уже пойманный тег, чтобы ловить несколько тегов в строке,
        # иначе будет ловиться только первый тег в строке
        number_of_tags -= 1  # проверка что все теги из строки скопированы
    return list_of_tags_from_strs


def find_and_append_tags_from_paragraphs():
    list_of_tags_from_paragraphs = []
    for paragraph in doc.paragraphs:    # весь текст документа делится на paragraphs, по сути paragraphs = абзацы
        tags = how_many_tags(paragraph.text)
        paragraph_text = paragraph.text
        list_of_tags_from_paragraphs = find_tag_in_text(tags, paragraph_text, list_of_tags_from_paragraphs)
    return list_of_tags_from_paragraphs


def find_and_append_tags_from_tables():
    list_of_tags_from_tables = []
    for table in doc.tables:    # все таблицы в документе делятся на таблицы, в них находим строки, в строках ячейки,
        for row in table.rows:  # а в ячейках текст
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    tags = how_many_tags(paragraph.text)
                    paragraph_text = paragraph.text
                    list_of_tags_from_tables = find_tag_in_text(tags, paragraph_text, list_of_tags_from_tables)
    return list_of_tags_from_tables


def making_dict_and_adding_keys(tags_from_paragraphs, tags_from_tables):
    list_of_tags = list(set(tags_from_paragraphs + tags_from_tables))  # суммируем списки, делаем из них множество,
    dict_of_tags = {}  # чтобы избавиться от повторения, потом снова приводим к списку
    for i in range(0, len(list_of_tags)): # присваиваем тегам значение 'value of tag' без {{ и }}
        dict_of_tags[list_of_tags[i]] = 'value of ' + list_of_tags[i]
    for value in dict_of_tags.items():
        new_value = input('Введите значение тега "{}": '.format(value[0]))  # присваиваем тегам свои значения
        if new_value != '':  # условие, чтобы пустые строки не попали в словарь, то есть если прожимать Enter на
            dict_of_tags[value[0]] = new_value  # input'е в значение будет отправляться 'value of tag'
    return dict_of_tags


def rewriting_tags_to_doc(dict):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for value in dict.items():
                if value[0] in run.text:  # run - послед-ность символов в едином форматировании символов, то есть
                    start = run.text.find('{{' + value[0])  # run это сущность внутри paragraphs которая объединяет
                    end = run.text.find(value[0] + '}}', start) # текст в абзаце с единым форматироватием
                    run.text = run.text[:start] + value[1] + run.text[end + 2 + len(value[0]):]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for value in dict.items():
                            if value[0] in run.text:  # Без run текст вставляется в документ без форматирования, которое
                                start = run.text.find('{{' + value[0])  # использовалось в документе
                                end = run.text.find(value[0] + '}}', start)
                                run.text = run.text[:start] + value[1] + run.text[end + 2 + len(value[0]):]


d = making_dict_and_adding_keys(find_and_append_tags_from_paragraphs(), find_and_append_tags_from_tables())
rewriting_tags_to_doc(d)
doc.save('test_updated.docx')
